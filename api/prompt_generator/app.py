#!/usr/bin/env python
"""
api/prompt_generator/app.py

FastAPI service for:
 - Generating Job Descriptions   (/generate_jd)
 - Generating Interview Questions (/generate_questions)
 - Improving an existing JD file  (/improve_jd)

All AI calls use Gemini REST API with API key.
"""

import os
import io
import logging
from typing import List
from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv
from jinja2 import Environment, FileSystemLoader, select_autoescape
import google.generativeai as genai
import PyPDF2
from api.auth.auth import router as auth_router

app = FastAPI(title="Prompt Generator (Gemini API)")

# Configure CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Adjust for production to restrict origins
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Mount auth endpoints
app.include_router(auth_router, prefix="/auth", tags=["auth"])

try:
    from docx import Document
except ImportError:
    def Document(file_obj):
        raise HTTPException(status_code=500, detail="DOCX processing not available. Install python-docx")

try:
    from PIL import Image
    import pytesseract
except ImportError:
    def pytesseract_image_to_string(img):
        raise HTTPException(status_code=500, detail="Image processing not available. Install pillow and pytesseract")

# Load environment
load_dotenv()

API_KEY = os.getenv("GEMINI_API_KEY")
CHAT_MODEL = os.getenv("GEMINI_CHAT_MODEL", "gemini-1.5-flash")

if not API_KEY:
    raise RuntimeError("GEMINI_API_KEY must be set in .env")

# Configure Gemini
genai.configure(api_key=API_KEY)

# Setup Jinja2 environment
HERE = os.path.dirname(__file__)
TEMPLATE_DIR = os.path.join(HERE, "templates")
env = Environment(
    loader=FileSystemLoader(TEMPLATE_DIR),
    autoescape=select_autoescape(["j2"]),
    trim_blocks=True,
    lstrip_blocks=True
)

# Request/response schemas
class RequestModel(BaseModel):
    title: str
    level: str
    department: str
    chunks: List[str]

class JDResponse(BaseModel):
    job_description: str

class QuestionsResponse(BaseModel):
    interview_questions: str

# Helper to call Gemini chat API
def call_gemini_chat(system: str, user: str, temperature: float) -> str:
    try:
        model = genai.GenerativeModel(CHAT_MODEL)
        
        # Combine system and user messages
        prompt = f"{system}\n\n{user}"
        
        response = model.generate_content(
            prompt,
            generation_config=genai.types.GenerationConfig(
                temperature=temperature,
                max_output_tokens=2048,
            )
        )
        
        return response.text.strip()
        
    except Exception as e:
        logging.error(f"Gemini API error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Gemini API error: {str(e)}")

# Helper to extract text from uploaded files
def extract_text_from_file(contents: bytes, filename: str) -> str:
    ext = filename.lower().rsplit('.', 1)[-1]
    
    def extract_pdf() -> str:
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(contents))
            return '\n'.join(page.extract_text() or '' for page in reader.pages)
        except Exception as e:
            logging.error(f"PDF processing error: {str(e)}", exc_info=True)
            raise HTTPException(status_code=500, detail=f"PDF processing error: {str(e)}")
    
    def extract_docx() -> str:
        try:
            doc = Document(io.BytesIO(contents))
            return '\n'.join(p.text for p in doc.paragraphs)
        except Exception as e:
            logging.error(f"DOCX processing error: {str(e)}", exc_info=True)
            raise HTTPException(status_code=500, detail=f"DOCX processing error: {str(e)}")
    
    def extract_image() -> str:
        try:
            return pytesseract.image_to_string(Image.open(io.BytesIO(contents)))
        except Exception as e:
            logging.error(f"Image processing error: {str(e)}", exc_info=True)
            raise HTTPException(status_code=500, detail=f"Image processing error: {str(e)}")
    
    def extract_txt() -> str:
        try:
            return contents.decode('utf-8')
        except UnicodeDecodeError:
            try:
                return contents.decode('latin-1')
            except Exception as e:
                logging.error(f"Text encoding error: {str(e)}", exc_info=True)
                raise HTTPException(status_code=500, detail=f"Text encoding error: {str(e)}")
    
    handlers = {
        'pdf': extract_pdf,
        'docx': extract_docx,
        'jpg': extract_image,
        'jpeg': extract_image,
        'png': extract_image,
        'txt': extract_txt,
    }
    
    handler = handlers.get(ext)
    if handler:
        return handler()
    else:
        logging.error(f"Unsupported file type: {ext}")
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}. Supported: PDF, DOCX, TXT, JPG/PNG")

# Endpoint: generate JD
@app.post("/generate_jd", response_model=JDResponse)
def generate_jd(req: RequestModel) -> JDResponse:
    template = env.get_template("jd_generation.j2")
    user_prompt = template.render(**req.dict(), chunks=req.chunks)
    system_msg = "You are a helpful assistant that crafts clear, concise job descriptions."
    job_desc = call_gemini_chat(system_msg, user_prompt, temperature=0.3)
    return JDResponse(job_description=job_desc)

# Endpoint: generate interview questions
@app.post("/generate_questions", response_model=QuestionsResponse)
def generate_questions(req: RequestModel) -> QuestionsResponse:
    template = env.get_template("interview_questions.j2")
    user_prompt = template.render(**req.dict(), chunks=req.chunks)
    system_msg = "You are an expert interviewer who produces behavioral and technical questions."
    questions = call_gemini_chat(system_msg, user_prompt, temperature=0.4)
    return QuestionsResponse(interview_questions=questions)

# Endpoint: improve existing JD file
@app.post("/improve_jd")
async def improve_jd(file: UploadFile = File(...)) -> dict:
    contents = await file.read()
    raw_text = extract_text_from_file(contents, file.filename)
    template = env.get_template("jd_improvement.j2")
    prompt = template.render(raw=raw_text)
    system_msg = "You are an AI assistant that improves and enhances job descriptions."
    improved = call_gemini_chat(system_msg, prompt, temperature=0.3)
    return {"improved_jd": improved}

# Health check
@app.get("/ping")
def ping() -> dict:
    return {"pong": True}
